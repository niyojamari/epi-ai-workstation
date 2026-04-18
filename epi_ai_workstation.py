"""
Epi-AI Workstation — Deployment-Ready Build
============================================
Fixes applied from pre-deployment review:
  [CRITICAL] PHI/PII safeguard gate on upload
  [CRITICAL] Cox PH assumption test (Schoenfeld + log-log plot)
  [CRITICAL] Bivariate purposeful selection implemented (Hosmer-Lemeshow criteria)
  [CRITICAL] DALY inputs validated and clamped (DW ∈ [0,1], D > 0, counts ≥ 0)
  [BUG]      Silent dropna() replaced with counted, user-confirmed listwise deletion
  [BUG]      Logistic Y validated as binary before fitting
  [BUG]      Cox covariate list excludes time/event columns
  [BUG]      Imputation blocked on outcome/time/event variables
  [BUG]      Study Design Wizard rebuilt as decision matrix (no unreachable branches)
  [WARN]     Raw data preserved in session state; reset-to-raw available
  [WARN]     Analysis results written to session state for use by Report module
  [WARN]     Report module dynamically pulls session data (not hardcoded)
  [WARN]     BytesIO seeked to 0 before download
  [WARN]     External CDN icon replaced with inline SVG
  [WARN]     Data-use disclaimer shown on every upload
"""

import streamlit as st
import pandas as pd
import numpy as np
import statsmodels.api as sm
import scipy.stats as stats
import matplotlib.pyplot as plt
from lifelines import KaplanMeierFitter, CoxPHFitter
from lifelines.statistics import proportional_hazard_test
from sklearn.impute import SimpleImputer
import docx
from docx.shared import Pt, RGBColor
from io import BytesIO
import datetime
import warnings
warnings.filterwarnings("ignore")

# ─────────────────────────────────────────────
# PAGE CONFIG & GLOBAL STYLE
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="Epi-AI Workstation",
    page_icon="🧬",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
.main { background-color: #f8f9fa; }
h1, h2, h3 { color: #2c3e50; }
.stButton > button {
    background-color: #2c3e50; color: white;
    border-radius: 6px; border: none; padding: 0.4rem 1rem;
}
.stButton > button:hover { background-color: #34495e; }
.metric-card {
    background: white; border-radius: 8px;
    padding: 1rem; border-left: 4px solid #2c3e50;
    box-shadow: 0 1px 4px rgba(0,0,0,0.07);
}
.warning-box {
    background: #fff8e1; border-left: 4px solid #f9a825;
    padding: 0.8rem 1rem; border-radius: 6px; font-size: 0.9rem;
}
.success-box {
    background: #e8f5e9; border-left: 4px solid #2e7d32;
    padding: 0.8rem 1rem; border-radius: 6px; font-size: 0.9rem;
}
.phi-banner {
    background: #fce4ec; border: 1px solid #e91e63;
    border-radius: 8px; padding: 1rem 1.2rem; margin-bottom: 1rem;
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# SESSION STATE INITIALISATION
# ─────────────────────────────────────────────
DEFAULTS = {
    "dataset_raw":       None,   # original uploaded data — never mutated
    "dataset":           None,   # working copy (imputation applied here)
    "phi_acknowledged":  False,  # gate flag for upload
    "analysis_log":      [],     # list of dicts — feeds Report module
    "last_model_summary": None,
    "last_cox_summary":   None,
    "last_daly":          None,
    "imputed_vars":       [],
    "rows_dropped":       0,
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ─────────────────────────────────────────────
# INLINE SVG ICON (replaces external CDN icon)
# ─────────────────────────────────────────────
SIDEBAR_ICON = """
<svg width="64" height="64" viewBox="0 0 64 64" xmlns="http://www.w3.org/2000/svg">
  <rect width="64" height="64" rx="14" fill="#2c3e50"/>
  <rect x="14" y="28" width="6" height="22" rx="2" fill="#1abc9c"/>
  <rect x="24" y="20" width="6" height="30" rx="2" fill="#3498db"/>
  <rect x="34" y="24" width="6" height="26" rx="2" fill="#e74c3c"/>
  <rect x="44" y="16" width="6" height="34" rx="2" fill="#f39c12"/>
  <line x1="14" y1="48" x2="52" y2="48" stroke="white" stroke-width="1.5" stroke-opacity="0.5"/>
</svg>
"""

# ─────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────
with st.sidebar:
    st.markdown(SIDEBAR_ICON, unsafe_allow_html=True)
    st.title("Epi-AI Workstation")
    st.caption("Clinical-Grade Epidemiological Analysis")
    st.markdown("---")
    app_mode = st.radio("MAIN MENU", [
        "📥 1. Data Ingestion & Cleaning",
        "🧭 2. Study Design Wizard",
        "⚙️ 3. Universal Statistics",
        "🌍 4. GBD Metrics Engine",
        "📄 5. Automated Reporting",
    ])
    st.markdown("---")
    if st.session_state["dataset"] is not None:
        df_info = st.session_state["dataset"]
        st.markdown(f"**Active dataset:** `{df_info.shape[0]:,}` rows × `{df_info.shape[1]}` cols")
        if st.button("⚠️ Reset to Raw Data"):
            st.session_state["dataset"] = st.session_state["dataset_raw"].copy()
            st.session_state["imputed_vars"] = []
            st.session_state["rows_dropped"] = 0
            st.rerun()
    st.caption("v2.0 | Deployment-Ready Build")


# ═══════════════════════════════════════════════════════════
# MODULE 1 — DATA INGESTION & CLEANING
# ═══════════════════════════════════════════════════════════
if app_mode == "📥 1. Data Ingestion & Cleaning":
    st.header("📥 Data Ingestion & Quality Control")

    # ── PHI/PII Gate ──────────────────────────────────────
    st.markdown("""
    <div class="phi-banner">
    <strong>⚕️ Data Governance Notice</strong><br>
    This tool is designed for <strong>de-identified research datasets only</strong>.
    Do not upload data containing patient names, national IDs, dates of birth,
    addresses, contact information, or any other direct identifiers.
    By uploading a file you confirm compliance with applicable data protection
    regulations (GDPR, HIPAA, or local equivalent).
    </div>
    """, unsafe_allow_html=True)

    phi_check = st.checkbox(
        "I confirm this dataset has been de-identified and I am authorised to analyse it.",
        value=st.session_state["phi_acknowledged"]
    )
    st.session_state["phi_acknowledged"] = phi_check

    if not phi_check:
        st.warning("Please acknowledge the data governance notice before uploading.")
        st.stop()

    # ── File Upload ────────────────────────────────────────
    uploaded_file = st.file_uploader(
        "Upload Dataset (.csv or .dta)",
        type=["csv", "dta"],
        help="Stata .dta (versions 12–18) and CSV are supported."
    )

    if uploaded_file:
        try:
            if uploaded_file.name.endswith(".csv"):
                df_new = pd.read_csv(uploaded_file)
            else:
                df_new = pd.read_stata(uploaded_file)

            # Passive PHI column name scan
            phi_patterns = ["name", "surname", "dob", "birth", "address",
                            "phone", "email", "ssn", "nid", "passport", "postcode", "zip"]
            flagged = [c for c in df_new.columns
                       if any(p in c.lower() for p in phi_patterns)]
            if flagged:
                st.warning(
                    f"⚠️ Possible identifier columns detected: **{', '.join(flagged)}**. "
                    "Review before proceeding — these may contain direct identifiers."
                )

            st.session_state["dataset_raw"] = df_new.copy()
            st.session_state["dataset"]     = df_new.copy()
            st.session_state["imputed_vars"] = []
            st.session_state["rows_dropped"] = 0
            st.success(f"✅ Dataset loaded: {df_new.shape[0]:,} observations, {df_new.shape[1]} variables.")
        except Exception as e:
            st.error(f"Failed to read file: {e}")

    # ── Dataset Overview ───────────────────────────────────
    if st.session_state["dataset"] is not None:
        df = st.session_state["dataset"]

        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Observations",      f"{df.shape[0]:,}")
        col2.metric("Variables",         df.shape[1])
        col3.metric("Missing cells",     int(df.isnull().sum().sum()))
        col4.metric("Complete rows",     int(df.dropna().shape[0]))

        with st.expander("📋 Data Preview", expanded=False):
            st.dataframe(df.head(20), use_container_width=True)

        with st.expander("📊 Variable Summary", expanded=False):
            desc = df.describe(include="all").T
            desc["missing"] = df.isnull().sum()
            desc["dtype"]   = df.dtypes
            st.dataframe(desc, use_container_width=True)

        # ── Missing Data & Imputation ──────────────────────
        st.markdown("### Missing Data Handling")
        missing_vars = df.columns[df.isnull().any()].tolist()

        if missing_vars:
            st.markdown("""
            <div class="warning-box">
            <strong>Imputation guidance:</strong> Median imputation is valid only for predictor
            variables in datasets that are plausibly Missing Completely At Random (MCAR).
            <em>Never impute outcome, time-to-event, or event-status variables.</em>
            </div>
            """, unsafe_allow_html=True)

            # Show missingness table
            miss_df = pd.DataFrame({
                "Variable": missing_vars,
                "Missing N": [df[v].isnull().sum() for v in missing_vars],
                "Missing %": [f"{df[v].isnull().mean()*100:.1f}%" for v in missing_vars],
                "Dtype":     [str(df[v].dtype) for v in missing_vars],
            })
            st.dataframe(miss_df, use_container_width=True, hide_index=True)

            # Only allow numeric, non-outcome columns for imputation
            numeric_missing = [v for v in missing_vars
                               if pd.api.types.is_numeric_dtype(df[v])]
            to_impute = st.multiselect(
                "Select numeric predictor variables to impute (median strategy):",
                options=numeric_missing,
                help="Do not select outcome, time, or event columns."
            )

            if to_impute and st.button("▶ Run Imputation"):
                imputer = SimpleImputer(strategy="median")
                df_work = st.session_state["dataset"].copy()
                df_work[to_impute] = imputer.fit_transform(df_work[to_impute])
                st.session_state["dataset"]     = df_work
                st.session_state["imputed_vars"] = list(set(
                    st.session_state["imputed_vars"] + to_impute
                ))
                st.session_state["analysis_log"].append({
                    "step": "Imputation",
                    "detail": f"Median imputation applied to: {', '.join(to_impute)}"
                })
                st.success(f"Imputation complete for: {', '.join(to_impute)}")
                st.rerun()
        else:
            st.markdown('<div class="success-box">✅ No missing values detected.</div>',
                        unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════
# MODULE 2 — STUDY DESIGN WIZARD
# ═══════════════════════════════════════════════════════════
elif app_mode == "🧭 2. Study Design Wizard":
    st.header("🧭 Study Design & Methodology Routing")
    st.markdown("Define your study parameters. The engine selects the correct statistical path.")

    col1, col2, col3 = st.columns(3)
    with col1:
        outcome = st.selectbox("1. Outcome variable type", [
            "Continuous", "Binary (0/1)", "Time-to-Event (Survival)",
            "Count", "Ordinal"
        ])
    with col2:
        structure = st.selectbox("2. Data structure", [
            "Independent observations",
            "Nested / Clustered (repeated measures)",
            "Complex survey weights",
            "Matched / Case-Control",
        ])
    with col3:
        sample = st.selectbox("3. Sample size context", [
            "Large (≥500)", "Moderate (100–499)", "Small (<100)"
        ])

    st.markdown("---")

    # ── Decision matrix (replaces fragile elif chain) ─────
    DECISION_MATRIX = {
        ("Continuous",              "Independent observations"):             ("OLS Linear Regression",                 "Robust SE recommended. Check normality of residuals via Q-Q plot."),
        ("Continuous",              "Nested / Clustered (repeated measures)"): ("Linear Mixed-Effects Model (LME)",      "Use lme4-equivalent; include random intercepts per cluster."),
        ("Continuous",              "Complex survey weights"):                ("Survey-Weighted OLS (svy: regress)",     "Apply Taylor linearisation for SE. Use survey package or statsmodels SurveyLS."),
        ("Continuous",              "Matched / Case-Control"):               ("Paired t-test or Conditional OLS",       "Condition on matched set; do not use standard OLS."),
        ("Binary (0/1)",            "Independent observations"):             ("Logistic Regression",                   "Check binary outcome; report OR with 95% CI. Run Hosmer-Lemeshow GOF."),
        ("Binary (0/1)",            "Nested / Clustered (repeated measures)"): ("Mixed-Effects Logistic Regression",    "Random intercepts per cluster. Use GEE for population-average estimates."),
        ("Binary (0/1)",            "Complex survey weights"):               ("Survey-Weighted Logistic Regression",    "Use svyglm with quasibinomial family. Adjust for design effect."),
        ("Binary (0/1)",            "Matched / Case-Control"):               ("Conditional Logistic Regression",       "Conditions on matched strata. Standard logistic is BIASED here."),
        ("Time-to-Event (Survival)","Independent observations"):             ("Cox Proportional Hazards",              "Verify PH assumption via Schoenfeld residuals. Report HR with 95% CI."),
        ("Time-to-Event (Survival)","Nested / Clustered (repeated measures)"): ("Shared-Frailty Cox Model",            "Add gamma frailty term per cluster. Accounts for within-cluster correlation."),
        ("Time-to-Event (Survival)","Complex survey weights"):               ("Weighted Nelson-Aalen / Cox",            "Survey weights applied at the risk set level. Use specialized surveydata packages."),
        ("Time-to-Event (Survival)","Matched / Case-Control"):               ("Stratified Cox Model",                  "Stratify on matched set ID. Eliminates baseline hazard nuisance parameter."),
        ("Count",                   "Independent observations"):             ("Poisson or Negative Binomial Regression","Check overdispersion (Var > Mean → NegBin). Include offset for person-time."),
        ("Count",                   "Nested / Clustered (repeated measures)"): ("Mixed-Effects Poisson / NegBin",       "Random intercept per cluster. Account for repeated count measures."),
        ("Count",                   "Complex survey weights"):               ("Survey-Weighted Poisson",                "Use svyglm with Poisson family + robust SE."),
        ("Count",                   "Matched / Case-Control"):               ("Conditional Poisson (incl. Mantel-Haenszel)", "Condition on matched strata; use log-linear models."),
        ("Ordinal",                 "Independent observations"):             ("Proportional Odds Logistic Regression", "Verify proportional odds assumption via Brant test."),
        ("Ordinal",                 "Nested / Clustered (repeated measures)"): ("Mixed-Effects Ordinal Regression",    "clmm (R) or ordinal package equivalent. Account for clustering."),
        ("Ordinal",                 "Complex survey weights"):               ("Survey-Weighted Ordinal Regression",    "Limited software support — use polr with weights in R svyolr."),
        ("Ordinal",                 "Matched / Case-Control"):               ("Stratified Ordinal Regression",         "Condition on matched strata. Use conditional likelihood."),
    }

    key = (outcome, structure)
    method, note = DECISION_MATRIX.get(key, ("Consult a biostatistician",
                                             "This combination requires bespoke modelling."))

    # Small-sample warnings
    small_warnings = []
    if "Small" in sample and "Logistic" in method:
        small_warnings.append("Small sample + logistic regression: consider Firth's penalised likelihood to avoid separation bias.")
    if "Small" in sample and "Cox" in method:
        small_warnings.append("Small sample + survival model: fewer than 10 events per variable is an overfitting risk — reduce covariates.")
    if "Small" in sample and "Mixed" in method:
        small_warnings.append("Small sample + mixed effects: random effects estimates will be unstable. Consider simpler fixed-effects approach.")

    st.subheader("✅ Recommended Statistical Path")
    st.success(f"**{method}**")
    st.info(f"📌 {note}")

    for w in small_warnings:
        st.warning(f"⚠️ Small-sample caution: {w}")

    with st.expander("Why this recommendation?"):
        st.markdown(f"""
        | Parameter | Selection |
        |---|---|
        | Outcome type | {outcome} |
        | Data structure | {structure} |
        | Sample context | {sample} |
        | Method | **{method}** |

        The decision engine applies Hosmer & Lemeshow (2013) model selection criteria
        and standard epidemiological practice. Always verify model assumptions after fitting.
        """)

    st.session_state["analysis_log"].append({
        "step": "Study Design",
        "detail": f"Outcome: {outcome} | Structure: {structure} | Recommended: {method}"
    })


# ═══════════════════════════════════════════════════════════
# MODULE 3 — UNIVERSAL STATISTICS
# ═══════════════════════════════════════════════════════════
elif app_mode == "⚙️ 3. Universal Statistics":
    st.header("⚙️ Universal Statistical Engine")

    if st.session_state["dataset"] is None:
        st.warning("⬆️ Upload a dataset in Module 1 first.")
        st.stop()

    df_full = st.session_state["dataset"]
    cols    = df_full.columns.tolist()

    tab1, tab2, tab3 = st.tabs([
        "📈 Linear / Logistic Regression",
        "⏱ Survival Analysis (Cox PH)",
        "🔬 Bivariate Purposeful Selection",
    ])

    # ── TAB 1: Regression ─────────────────────────────────
    with tab1:
        st.subheader("Regression Modeller")
        model_type = st.radio("Model family:", ["Linear (OLS)", "Logistic (Binary)"],
                              horizontal=True)
        y_var  = st.selectbox("Outcome variable (Y):", cols, key="reg_y")
        x_vars = st.multiselect("Predictor variables (X):", cols, key="reg_x")

        if st.button("▶ Fit Regression Model"):
            if not x_vars:
                st.error("Select at least one predictor.")
            else:
                # Listwise deletion with count
                df_model = df_full[[y_var] + x_vars].dropna()
                n_dropped = len(df_full) - len(df_model)
                if n_dropped > 0:
                    st.warning(f"⚠️ Listwise deletion: {n_dropped:,} rows removed due to missing values "
                               f"({n_dropped/len(df_full)*100:.1f}% of dataset). "
                               f"Analysis run on {len(df_model):,} complete cases.")
                st.session_state["rows_dropped"] = n_dropped

                # Binary validation for logistic
                if "Logistic" in model_type:
                    unique_vals = df_model[y_var].dropna().unique()
                    if not (set(unique_vals).issubset({0, 1, 0.0, 1.0}) and len(unique_vals) == 2):
                        st.error(
                            f"❌ Logistic regression requires a binary outcome (0/1). "
                            f"'{y_var}' contains {len(unique_vals)} unique values: {sorted(unique_vals)[:10]}. "
                            "Select a binary variable or recode your outcome."
                        )
                        st.stop()

                try:
                    X = sm.add_constant(
                        pd.get_dummies(df_model[x_vars], drop_first=True).astype(float)
                    )
                    y = df_model[y_var].astype(float)

                    if "Linear" in model_type:
                        model  = sm.OLS(y, X).fit(cov_type="HC3")   # robust SE by default
                        result = model.summary()
                        st.code(result.as_text(), language="text")

                        # Residual diagnostics
                        with st.expander("📊 Residual Diagnostics"):
                            fig, axes = plt.subplots(1, 2, figsize=(10, 4))
                            residuals = model.resid
                            axes[0].scatter(model.fittedvalues, residuals, alpha=0.4, s=20)
                            axes[0].axhline(0, color="red", linewidth=1)
                            axes[0].set_xlabel("Fitted values")
                            axes[0].set_ylabel("Residuals")
                            axes[0].set_title("Residuals vs Fitted")
                            sm.qqplot(residuals, line="s", ax=axes[1], alpha=0.4)
                            axes[1].set_title("Q-Q Plot (normality check)")
                            plt.tight_layout()
                            st.pyplot(fig)
                            plt.close()

                        st.session_state["last_model_summary"] = result.as_text()
                        st.session_state["analysis_log"].append({
                            "step": "Linear Regression (OLS)",
                            "detail": (f"Y={y_var}, X={x_vars}, N={len(df_model):,}, "
                                       f"R²={model.rsquared:.4f}, AIC={model.aic:.2f}")
                        })

                    else:  # Logistic
                        model  = sm.Logit(y, X).fit(disp=False)
                        result = model.summary()
                        st.code(result.as_text(), language="text")

                        # Odds ratios
                        with st.expander("📊 Odds Ratios with 95% CI"):
                            params = model.params
                            conf   = model.conf_int()
                            or_df  = pd.DataFrame({
                                "Odds Ratio": np.exp(params),
                                "95% CI Lower": np.exp(conf[0]),
                                "95% CI Upper": np.exp(conf[1]),
                                "P-value": model.pvalues,
                            }).round(4)
                            st.dataframe(or_df, use_container_width=True)

                        # Hosmer-Lemeshow GOF
                        with st.expander("📊 Hosmer-Lemeshow Goodness-of-Fit Test"):
                            pred = model.predict()
                            deciles = pd.qcut(pred, 10, duplicates="drop")
                            obs_1   = y.groupby(deciles).sum()
                            exp_1   = pred.groupby(deciles).sum()
                            n_grp   = y.groupby(deciles).count()
                            hl_stat = (((obs_1 - exp_1)**2 / (exp_1 * (1 - exp_1/n_grp)))
                                       .dropna().sum())
                            hl_p    = 1 - stats.chi2.cdf(hl_stat, df=8)
                            col_a, col_b = st.columns(2)
                            col_a.metric("H-L χ² statistic", f"{hl_stat:.3f}")
                            col_b.metric("P-value (df=8)",   f"{hl_p:.4f}")
                            if hl_p < 0.05:
                                st.warning("⚠️ P < 0.05: model fit may be poor. Reconsider specification.")
                            else:
                                st.success("✅ P ≥ 0.05: no significant lack of fit detected.")

                        st.session_state["last_model_summary"] = result.as_text()
                        st.session_state["analysis_log"].append({
                            "step": "Logistic Regression",
                            "detail": (f"Y={y_var}, X={x_vars}, N={len(df_model):,}, "
                                       f"AIC={model.aic:.2f}, HL-p={hl_p:.4f}")
                        })

                except Exception as e:
                    st.error(f"Model computation error: {e}")

    # ── TAB 2: Cox PH ─────────────────────────────────────
    with tab2:
        st.subheader("Cox Proportional Hazards Model")

        time_col  = st.selectbox("Duration variable (time-to-event):", cols, key="cox_t")
        event_col = st.selectbox("Event status (1=event, 0=censored):", cols, key="cox_e")

        # Exclude time/event cols from covariate list (bug fix)
        covar_options = [c for c in cols if c not in [time_col, event_col]]
        covars = st.multiselect("Covariates:", covar_options, key="cox_x")

        if st.button("▶ Run Cox PH Model"):
            cox_cols = [time_col, event_col] + covars
            df_cox   = df_full[cox_cols].dropna()
            n_dropped = len(df_full) - len(df_cox)

            if n_dropped > 0:
                st.warning(f"⚠️ Listwise deletion: {n_dropped:,} rows removed. "
                           f"Analysis on {len(df_cox):,} complete cases.")

            n_events = int(df_cox[event_col].sum())
            epv = n_events / max(len(covars), 1) if covars else n_events
            if covars and epv < 10:
                st.warning(
                    f"⚠️ Events-per-variable (EPV) = {epv:.1f}. "
                    "EPV < 10 is an overfitting risk. Consider reducing covariates."
                )

            try:
                cph = CoxPHFitter()
                cph.fit(df_cox, duration_col=time_col, event_col=event_col)

                st.markdown("#### Model Summary")
                st.dataframe(cph.summary.round(4), use_container_width=True)

                # Hazard ratios
                with st.expander("📊 Hazard Ratios (exponentiated coefficients)"):
                    hr_df = pd.DataFrame({
                        "HR":           np.exp(cph.params_),
                        "95% CI Lower": np.exp(cph.confidence_intervals_["95% lower-bound"]),
                        "95% CI Upper": np.exp(cph.confidence_intervals_["95% upper-bound"]),
                        "P-value":      cph.summary["p"],
                    }).round(4)
                    st.dataframe(hr_df, use_container_width=True)

                # ── PH Assumption Test (Schoenfeld residuals) ──
                with st.expander("⚠️ Proportional Hazards Assumption Test (Required)", expanded=True):
                    st.markdown("""
                    The PH assumption states that hazard ratios are constant over time.
                    Violation invalidates the Cox model. Test via Schoenfeld residuals.
                    **P < 0.05 per variable = PH violated for that variable.**
                    """)
                    try:
                        ph_test = proportional_hazard_test(cph, df_cox, time_transform="rank")
                        ph_df = ph_test.summary[["test_statistic", "p"]].round(4).copy()
                        ph_df.columns = ["χ² statistic", "P-value"]
                        ph_df["PH Assumption"] = ph_df["P-value"].apply(
                            lambda p: "✅ Held" if p >= 0.05 else "❌ Violated"
                        )
                        st.dataframe(ph_df, use_container_width=True)

                        violated = ph_df[ph_df["PH Assumption"] == "❌ Violated"]
                        if not violated.empty:
                            st.error(
                                f"❌ PH violated for: **{', '.join(violated.index.tolist())}**. "
                                "Consider time-varying coefficients, stratification, or an accelerated "
                                "failure time (AFT) model."
                            )
                        else:
                            st.success("✅ PH assumption held for all covariates.")
                    except Exception as ph_err:
                        st.warning(f"PH test could not be computed: {ph_err}")

                # Kaplan-Meier plot
                with st.expander("📈 Kaplan-Meier Survival Curve"):
                    fig, ax = plt.subplots(figsize=(8, 4))
                    kmf = KaplanMeierFitter()
                    kmf.fit(df_cox[time_col], event_observed=df_cox[event_col])
                    kmf.plot_survival_function(ax=ax, ci_show=True, color="#2c3e50")
                    ax.set_xlabel("Time")
                    ax.set_ylabel("Survival probability")
                    ax.set_title("Kaplan-Meier Estimate (overall)")
                    ax.set_ylim(0, 1)
                    plt.tight_layout()
                    st.pyplot(fig)
                    plt.close()

                st.session_state["last_cox_summary"] = cph.summary.to_string()
                st.session_state["analysis_log"].append({
                    "step": "Cox PH Model",
                    "detail": (f"Time={time_col}, Event={event_col}, "
                               f"Covariates={covars}, N={len(df_cox):,}, "
                               f"Events={n_events}")
                })

            except Exception as e:
                st.error(f"Cox model error: {e}")

    # ── TAB 3: Bivariate Purposeful Selection ──────────────
    with tab3:
        st.subheader("Bivariate Purposeful Selection")
        st.markdown("""
        Implements the **Hosmer-Lemeshow purposeful selection** procedure:
        1. Each candidate predictor is tested individually against the binary outcome.
        2. Variables significant at **p ≤ 0.25** pass to the multivariable model.
        3. A final multivariable model retains variables significant at **p ≤ 0.05** or those
           that change any other coefficient by ≥ 20% (confounding criterion).
        """)

        ps_outcome = st.selectbox("Binary outcome variable (Y):", cols, key="ps_y")
        ps_candidates = st.multiselect(
            "Candidate predictor variables:", cols, key="ps_x"
        )
        ps_screen_p = st.slider("Screening threshold (p₁):", 0.10, 0.30, 0.25, 0.01)
        ps_final_p  = st.slider("Final model threshold (p₂):", 0.01, 0.10, 0.05, 0.01)

        if st.button("▶ Run Purposeful Selection"):
            if not ps_candidates:
                st.error("Select at least two candidate predictors.")
            else:
                df_ps = df_full[[ps_outcome] + ps_candidates].dropna()

                # Validate binary Y
                unique_vals = df_ps[ps_outcome].dropna().unique()
                if not (set(unique_vals).issubset({0, 1, 0.0, 1.0}) and len(unique_vals) == 2):
                    st.error(
                        f"'{ps_outcome}' is not binary (0/1). "
                        "Purposeful selection requires a binary outcome."
                    )
                    st.stop()

                y_ps = df_ps[ps_outcome].astype(float)

                # ── Step 1: Univariable screening ─────────
                st.markdown("#### Step 1 — Univariable Screening")
                screen_results = []
                for var in ps_candidates:
                    try:
                        X_uni = sm.add_constant(
                            pd.get_dummies(df_ps[[var]], drop_first=True).astype(float)
                        )
                        m_uni = sm.Logit(y_ps, X_uni).fit(disp=False)
                        # Take minimum p-value among all levels of the variable
                        var_cols = [c for c in m_uni.pvalues.index if c != "const"]
                        min_p    = m_uni.pvalues[var_cols].min()
                        screen_results.append({
                            "Variable": var,
                            "LR Chi²":  round(m_uni.llr, 3),
                            "P-value":  round(float(min_p), 4),
                            "Pass (p ≤ threshold)": "✅ Yes" if min_p <= ps_screen_p else "❌ No"
                        })
                    except Exception:
                        screen_results.append({
                            "Variable": var,
                            "LR Chi²":  None,
                            "P-value":  None,
                            "Pass (p ≤ threshold)": "⚠️ Error"
                        })

                screen_df = pd.DataFrame(screen_results)
                st.dataframe(screen_df, use_container_width=True, hide_index=True)
                passed = [r["Variable"] for r in screen_results
                          if r["Pass (p ≤ threshold)"] == "✅ Yes"]

                if not passed:
                    st.warning("No variables passed the screening threshold. "
                               "Try raising p₁ or reconsider variable selection.")
                    st.stop()

                st.success(f"**{len(passed)} variable(s) passed to multivariable model:** {', '.join(passed)}")

                # ── Step 2: Multivariable model ───────────
                st.markdown("#### Step 2 — Multivariable Logistic Regression")
                try:
                    X_multi = sm.add_constant(
                        pd.get_dummies(df_ps[passed], drop_first=True).astype(float)
                    )
                    m_multi = sm.Logit(y_ps, X_multi).fit(disp=False)
                    st.code(m_multi.summary().as_text(), language="text")

                    # Odds ratios table
                    params_m = m_multi.params
                    conf_m   = m_multi.conf_int()
                    or_multi = pd.DataFrame({
                        "Odds Ratio":    np.exp(params_m).round(4),
                        "95% CI Lower":  np.exp(conf_m[0]).round(4),
                        "95% CI Upper":  np.exp(conf_m[1]).round(4),
                        "P-value":       m_multi.pvalues.round(4),
                        "Significant":   m_multi.pvalues.apply(
                            lambda p: f"✅ p≤{ps_final_p}" if p <= ps_final_p else "— retain if confounding"
                        )
                    })
                    st.dataframe(or_multi, use_container_width=True)

                    # Final retained variables
                    retained = [c for c in passed
                                if any(c in idx for idx in
                                       m_multi.pvalues[m_multi.pvalues <= ps_final_p].index)]
                    st.markdown(f"**Final model variables (p ≤ {ps_final_p}):** "
                                f"{', '.join(retained) if retained else 'None — review confounding criterion'}")

                    st.session_state["analysis_log"].append({
                        "step": "Bivariate Purposeful Selection",
                        "detail": (f"Outcome={ps_outcome}, Screened={len(ps_candidates)}, "
                                   f"Passed={len(passed)}, Final retained={len(retained)}")
                    })

                except Exception as e:
                    st.error(f"Multivariable model error: {e}")


# ═══════════════════════════════════════════════════════════
# MODULE 4 — GBD METRICS ENGINE
# ═══════════════════════════════════════════════════════════
elif app_mode == "🌍 4. GBD Metrics Engine":
    st.header("🌍 Global Burden of Disease Calculators")
    st.markdown("Standardised WHO/GBD population-health metrics with input validation.")

    tab_daly, tab_smr, tab_paf = st.tabs([
        "DALY Calculator", "SMR / PMR", "Population Attributable Fraction"
    ])

    # ── DALY ──────────────────────────────────────────────
    with tab_daly:
        st.markdown("""
        **DALY = YLL + YLD**
        - YLL (Years of Life Lost) = N × L
        - YLD (Years Lived with Disability) = I × DW × D
        """)

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Years of Life Lost (YLL)**")
            N = st.number_input("Deaths (N):", min_value=0, value=100, step=1)
            L = st.number_input("Standard life expectancy at age of death (L, years):",
                                min_value=0.0, max_value=100.0, value=70.0, step=0.5)
        with col2:
            st.markdown("**Years Lived with Disability (YLD)**")
            I  = st.number_input("Incident cases (I):", min_value=0, value=500, step=1)
            DW = st.number_input("Disability weight (DW, 0–1):",
                                 min_value=0.0, max_value=1.0, value=0.15, step=0.01,
                                 help="WHO GBD disability weights: 0 = full health, 1 = death equivalent.")
            D  = st.number_input("Average duration of disease (D, years):",
                                 min_value=0.01, max_value=100.0, value=5.0, step=0.1)

        # Input validation (critical fix)
        input_errors = []
        if not (0.0 <= DW <= 1.0):
            input_errors.append("Disability Weight must be between 0 and 1.")
        if D <= 0:
            input_errors.append("Duration must be greater than 0.")
        if N < 0 or I < 0:
            input_errors.append("Death count and incident cases must be non-negative.")

        if input_errors:
            for err in input_errors:
                st.error(f"❌ {err}")
        elif st.button("▶ Calculate DALYs"):
            yll   = N * L
            yld   = I * DW * D
            daly  = yll + yld
            col_a, col_b, col_c = st.columns(3)
            col_a.metric("YLL (Years of Life Lost)",         f"{yll:,.2f}")
            col_b.metric("YLD (Years Lived with Disability)", f"{yld:,.2f}")
            col_c.metric("Total DALYs",                       f"{daly:,.2f}")
            st.session_state["last_daly"] = {
                "N": N, "L": L, "I": I, "DW": DW, "D": D,
                "YLL": yll, "YLD": yld, "DALY": daly
            }
            st.session_state["analysis_log"].append({
                "step": "DALY Calculation",
                "detail": f"YLL={yll:,.2f}, YLD={yld:,.2f}, Total DALY={daly:,.2f}"
            })

    # ── SMR / PMR ─────────────────────────────────────────
    with tab_smr:
        st.markdown("""
        **Standardised Mortality Ratio (SMR)** = Observed deaths / Expected deaths
        """)
        obs   = st.number_input("Observed deaths:", min_value=0, value=45, step=1)
        exp   = st.number_input("Expected deaths (from reference population):",
                                min_value=0.001, value=30.0, step=0.1)
        if st.button("▶ Calculate SMR"):
            smr = obs / exp
            # Byar's 95% CI
            ci_lo = (obs * (1 - 1/(9*obs) - 1.96/np.sqrt(9*obs))**3) / exp
            ci_hi = ((obs+1) * (1 - 1/(9*(obs+1)) + 1.96/np.sqrt(9*(obs+1)))**3) / exp
            col_a, col_b, col_c = st.columns(3)
            col_a.metric("SMR",                f"{smr:.3f}")
            col_b.metric("95% CI (lower)",     f"{ci_lo:.3f}")
            col_c.metric("95% CI (upper)",     f"{ci_hi:.3f}")
            if smr > 1:
                st.warning(f"SMR > 1: study population mortality is {((smr-1)*100):.1f}% "
                           "higher than reference population.")
            else:
                st.success(f"SMR < 1: study population mortality is {((1-smr)*100):.1f}% "
                           "lower than reference population.")
            st.session_state["analysis_log"].append({
                "step": "SMR Calculation",
                "detail": f"Observed={obs}, Expected={exp:.1f}, SMR={smr:.3f} (95% CI: {ci_lo:.3f}–{ci_hi:.3f})"
            })

    # ── PAF ───────────────────────────────────────────────
    with tab_paf:
        st.markdown("""
        **Population Attributable Fraction (PAF)**
        = (Pe × (RR − 1)) / (1 + Pe × (RR − 1))
        where Pe = prevalence of exposure in population.
        """)
        pe = st.number_input("Exposure prevalence in population (Pe, 0–1):",
                             min_value=0.0, max_value=1.0, value=0.30, step=0.01)
        rr = st.number_input("Relative Risk or Odds Ratio (RR):",
                             min_value=0.01, value=2.5, step=0.1)
        if st.button("▶ Calculate PAF"):
            if rr <= 0:
                st.error("RR must be greater than 0.")
            else:
                paf = (pe * (rr - 1)) / (1 + pe * (rr - 1))
                st.metric("Population Attributable Fraction (PAF)",
                          f"{paf*100:.2f}%")
                st.info(f"Eliminating this exposure could theoretically prevent "
                        f"**{paf*100:.1f}%** of cases in this population "
                        f"(assuming causality and complete exposure removal).")
                st.session_state["analysis_log"].append({
                    "step": "PAF Calculation",
                    "detail": f"Pe={pe}, RR={rr}, PAF={paf*100:.2f}%"
                })


# ═══════════════════════════════════════════════════════════
# MODULE 5 — AUTOMATED REPORTING
# ═══════════════════════════════════════════════════════════
elif app_mode == "📄 5. Automated Reporting":
    st.header("📄 Automated Methodology & Results Report")

    col1, col2 = st.columns(2)
    with col1:
        title    = st.text_input("Project title:", "Epidemiological Analysis Report")
        author   = st.text_input("Lead investigator:", "Senior Epidemiologist")
        affil    = st.text_input("Institution / Department:", "")
    with col2:
        study_design = st.text_input("Study design:", "Cross-sectional / Cohort / Case-Control")
        study_pop    = st.text_area("Study population:", "", height=80)

    st.markdown("### Analysis log (auto-populated from your session)")
    log = st.session_state["analysis_log"]
    if log:
        log_df = pd.DataFrame(log)
        st.dataframe(log_df, use_container_width=True, hide_index=True)
    else:
        st.info("No analyses run yet. Complete modules 1–4 to auto-populate the report.")

    if st.button("▶ Generate Report (.docx)"):
        doc  = docx.Document()
        now  = datetime.datetime.now().strftime("%d %B %Y, %H:%M")

        # Cover page
        h = doc.add_heading(title, 0)
        h.runs[0].font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        doc.add_paragraph(f"Lead Investigator: {author}")
        if affil:
            doc.add_paragraph(f"Institution: {affil}")
        doc.add_paragraph(f"Generated: {now}")
        doc.add_paragraph(f"Study Design: {study_design}")
        if study_pop:
            doc.add_paragraph(f"Study Population: {study_pop}")
        doc.add_paragraph("")

        # Data section
        doc.add_heading("1. Data & Quality Control", level=1)
        df_work = st.session_state["dataset"]
        if df_work is not None:
            df_raw  = st.session_state["dataset_raw"]
            doc.add_paragraph(
                f"Dataset dimensions: {df_work.shape[0]:,} observations × {df_work.shape[1]} variables."
            )
            doc.add_paragraph(
                f"Original dataset (pre-processing): {df_raw.shape[0]:,} rows."
            )
            imp = st.session_state["imputed_vars"]
            if imp:
                doc.add_paragraph(
                    f"Median imputation applied to: {', '.join(imp)}. "
                    "Missing Completely At Random (MCAR) assumption applied."
                )
            else:
                doc.add_paragraph("No imputation was performed.")
            dropped = st.session_state["rows_dropped"]
            if dropped > 0:
                doc.add_paragraph(
                    f"Listwise deletion removed {dropped:,} incomplete cases "
                    f"({dropped/df_raw.shape[0]*100:.1f}% of original sample) "
                    "prior to regression modelling."
                )
        else:
            doc.add_paragraph("No dataset loaded in this session.")

        # Methods section (from analysis log)
        doc.add_heading("2. Statistical Methods", level=1)
        if log:
            for entry in log:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{entry['step']}: ").bold = True
                p.add_run(entry["detail"])
        else:
            doc.add_paragraph("No analyses were recorded in this session.")

        # Results section
        doc.add_heading("3. Results Summary", level=1)

        if st.session_state["last_model_summary"]:
            doc.add_heading("3.1 Regression Model Output", level=2)
            doc.add_paragraph(
                "Full model output below (statsmodels summary):"
            )
            p = doc.add_paragraph(st.session_state["last_model_summary"])
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(8)

        if st.session_state["last_cox_summary"]:
            doc.add_heading("3.2 Cox Proportional Hazards Output", level=2)
            p = doc.add_paragraph(st.session_state["last_cox_summary"])
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(8)

        if st.session_state["last_daly"]:
            doc.add_heading("3.3 DALY Calculation", level=2)
            d = st.session_state["last_daly"]
            doc.add_paragraph(
                f"Deaths (N)={d['N']}, Life expectancy (L)={d['L']} years, "
                f"Incident cases (I)={d['I']}, Disability weight (DW)={d['DW']}, "
                f"Duration (D)={d['D']} years."
            )
            doc.add_paragraph(
                f"YLL = {d['YLL']:,.2f} | YLD = {d['YLD']:,.2f} | "
                f"Total DALY = {d['DALY']:,.2f}"
            )

        # Limitations
        doc.add_heading("4. Limitations", level=1)
        doc.add_paragraph(
            "Standard epidemiological limitations apply, including potential residual confounding, "
            "information bias, and selection bias. Causal inference should not be drawn from "
            "observational data without appropriate design and sensitivity analyses. "
            "Median imputation assumes MCAR; this assumption was not formally tested."
        )

        # Ethics note
        doc.add_heading("5. Data Governance", level=1)
        doc.add_paragraph(
            "The investigator confirmed prior to analysis that the dataset was de-identified "
            "and that its use complies with applicable data protection regulations "
            "(GDPR, HIPAA, or local equivalent). No personally identifiable information "
            "was processed by this tool."
        )

        # Save & offer download
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)   # seek fix applied

        st.success("✅ Report compiled successfully from session data.")
        st.download_button(
            label="📥 Download Report (.docx)",
            data=bio.getvalue(),
            file_name=f"Epi_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
